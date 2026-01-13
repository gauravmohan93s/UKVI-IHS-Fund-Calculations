from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

path = r"docs/UK_Visa_Calculator_User_Guide.docx"
logo_path = r"public/assets/kc_logo.png"


doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
style.font.size = Pt(11)

# Header with logo
if os.path.exists(logo_path):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(logo_path, width=Inches(1.2))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Title
p = doc.add_paragraph()
run = p.add_run("KC Overseas Education Pvt Ltd\nUK Visa Funds & IHS Calculator – User Guide")
run.bold = True
run.font.size = Pt(16)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

meta = doc.add_paragraph("Version: v2.3.0\nAudience: Counselors / Operations / Admissions")


def add_heading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)


def add_bullets(items):
    for i in items:
        doc.add_paragraph(i, style='List Bullet')


def add_placeholder(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)

add_heading("1) Purpose")
add_bullets([
    "Calculate IHS, funds required, funds available, and generate client-ready PDFs.",
    "Uses cached data for fast response and daily student data sync from Google Sheets.",
])

add_heading("2) Access & Security")
add_bullets([
    "Open the app URL shared internally.",
    "Enter the access code when prompted (do not share externally).",
    "Access code is required for all API calls; unauthorized users are blocked.",
])

add_heading("3) Quick Use – IHS Calculation")
add_bullets([
    "Go to Tab 1 (Course & application).",
    "Enter Course start and end dates.",
    "IHS total appears immediately in the IHS box.",
])
add_placeholder("[Insert Screenshot: IHS quick calculation box]")

add_heading("4) Quick Use – Funds Required")
add_bullets([
    "Enter Tuition total, Tuition paid, Scholarship, Dependants, Buffer.",
    "Set Visa application date (recommended for 31-day checks).",
    "Results show Funds Required in GBP and display currency.",
])
add_placeholder("[Insert Screenshot: Fees & funds input section]")

add_heading("5) Daily Workflow (Full Calculation)")
add_bullets([
    "Search student by acknowledgement number or name.",
    "Select counselor and verify auto-filled fields.",
    "Add funds rows (bank/FD/loan) with dates and amounts.",
    "Go to Results and review eligibility + warnings.",
    "Download PDF when ready.",
])
add_placeholder("[Insert Screenshot: Funds available breakdown + results]")

add_heading("6) Currency Handling")
add_bullets([
    "Display currency can be changed for client communication.",
    "If FX is unavailable, the system falls back to GBP-only display.",
    "Manual FX overrides can be used if live FX fails.",
])
add_placeholder("[Insert Screenshot: FX settings panel]")

add_heading("7) Data Updates")
add_bullets([
    "Students: Auto-synced daily from Google Sheets (11:00 AM IST).",
    "Counselors: Read from local CSV (update as required).",
    "Country currency: Read from local JSON (update when new countries appear).",
    "UKVI config: Update if GOV.UK policy changes.",
])

add_heading("8) Sync Status (Visibility)")
add_bullets([
    "The app shows the last successful student data update time in IST.",
    "If a sync fails, cached data remains available (no data loss).",
])

add_heading("9) Troubleshooting")
add_bullets([
    "FX error or missing conversions: retry or use manual FX override.",
    "No student results: confirm sheet is published and synced.",
    "PDF issues: re-check dates and required fields, then regenerate.",
])

add_heading("10) Support")
add_bullets([
    "For issues or access: Contact the internal operations lead.",
])

os.makedirs(os.path.dirname(path), exist_ok=True)
doc.save(path)
print(path)
